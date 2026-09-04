import os
import shutil
import tempfile
import uuid
import subprocess
import zipfile
import asyncio
import logging
from typing import List
from fastapi import FastAPI, UploadFile, File, Form, BackgroundTasks, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
startup_error = None
try:
    from PIL import Image
except Exception as e:
    Image = None
    startup_error = f"PIL error: {e}"

try:
    from docx import Document
except Exception as e:
    Document = None
    startup_error = f"docx error: {e}"

try:
    from pypdf import PdfReader, PdfWriter
except Exception as e:
    PdfReader = None
    PdfWriter = None
    startup_error = f"pypdf error: {e}"

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    reportlab_available = True
except Exception as e:
    reportlab_available = False

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Online File Converter API",
    description="Backend API for completely free All-in-One Online File Converter",
    version="1.0.0"
)


# Background task to clean up the temporary directory after a delay
async def cleanup_temp_dir(dir_path: str, delay: int = 180):
    """
    Asynchronously deletes the temporary directory and all its contents after a delay (default 3 minutes).
    This delay ensures the client has started and finished downloading the file.
    """
    await asyncio.sleep(delay)
    try:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
            logger.info(f"Successfully cleaned up temporary directory: {dir_path}")
    except Exception as e:
        logger.error(f"Error while cleaning up directory {dir_path}: {str(e)}")

# Ensure /tmp folder exists (FastAPI temp directory fallback)
temp_base_dir = os.path.join(tempfile.gettempdir(), "file_converter")
os.makedirs(temp_base_dir, exist_ok=True)

def create_request_temp_dir() -> str:
    """Creates a unique directory for the lifetime of a conversion request."""
    return tempfile.mkdtemp(dir=temp_base_dir, prefix="req_")

# 1. PDF to Word (.docx)
@app.post("/api/pdf-to-word")
async def pdf_to_word(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported for this operation.")
    
    temp_dir = create_request_temp_dir()
    input_path = os.path.join(temp_dir, "input.pdf")
    output_path = os.path.join(temp_dir, "converted.docx")
    
    try:
        # Save uploaded PDF file
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Run conversion using pure Python pypdf + docx
        reader = PdfReader(input_path)
        doc = Document()
        for idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                for line in text.splitlines():
                    clean_line = line.strip()
                    if clean_line:
                        doc.add_paragraph(clean_line)
            if idx < len(reader.pages) - 1:
                doc.add_page_break()
        doc.save(output_path)
        
        if not os.path.exists(output_path):
            raise Exception("Word (.docx) belgesi oluşturulamadı.")
        
        # Schedule cleanup task (3 minutes delay)
        background_tasks.add_task(cleanup_temp_dir, temp_dir)
        
        return FileResponse(
            path=output_path,
            filename=f"{os.path.splitext(file.filename)[0]}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"PDF to Word conversion failed: {str(e)}")
        # If it failed immediately, clean up the directory right away
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

# 2. Word (.docx) to PDF
@app.post("/api/word-to-pdf")
async def word_to_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    if not file.filename.lower().endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="Only Word documents (.doc, .docx) are supported.")
    
    temp_dir = create_request_temp_dir()
    input_path = os.path.join(temp_dir, file.filename)
    
    try:
        # Save uploaded docx file
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        base_name = os.path.splitext(file.filename)[0]
        output_path = os.path.join(temp_dir, f"{base_name}.pdf")
        
        doc = Document(input_path)
        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=40,
            leftMargin=40,
            topMargin=40,
            bottomMargin=40
        )
        styles = getSampleStyleSheet()
        story = []
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                style_name = getattr(p.style, 'name', '')
                if style_name.startswith('Heading 1'):
                    story.append(Paragraph(f"<b><font size=16>{safe_text}</font></b>", styles['Heading1']))
                elif style_name.startswith('Heading 2'):
                    story.append(Paragraph(f"<b><font size=14>{safe_text}</font></b>", styles['Heading2']))
                elif style_name.startswith('Heading 3'):
                    story.append(Paragraph(f"<b><font size=12>{safe_text}</font></b>", styles['Heading3']))
                else:
                    story.append(Paragraph(safe_text, styles['Normal']))
                story.append(Spacer(1, 6))
                
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [Paragraph(cell.text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), styles['Normal']) for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f1f5f9')),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                    ('TOPPADDING', (0,0), (-1,-1), 4),
                ]))
                story.append(t)
                story.append(Spacer(1, 10))
                
        if not story:
            story.append(Paragraph("(Boş Belge)", styles['Normal']))
            
        pdf_doc.build(story)
        
        if not os.path.exists(output_path):
            raise Exception("PDF belgesi oluşturulamadı.")
        
        # Schedule cleanup
        background_tasks.add_task(cleanup_temp_dir, temp_dir)
        
        return FileResponse(
            path=output_path,
            filename=f"{base_name}.pdf",
            media_type="application/pdf"
        )
    except Exception as e:
        logger.error(f"Word to PDF conversion failed: {str(e)}")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

# 3. PDF to Image (JPG/PNG) - Returns a ZIP file
@app.post("/api/pdf-to-image")
async def pdf_to_image(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    img_format: str = Form("png")  # png or jpg
):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
    
    img_format = img_format.lower()
    if img_format not in ["png", "jpg", "jpeg"]:
        raise HTTPException(status_code=400, detail="Invalid image format. Supported formats are png and jpg.")
    
    temp_dir = create_request_temp_dir()
    input_path = os.path.join(temp_dir, "input.pdf")
    zip_filename = f"{os.path.splitext(file.filename)[0]}_images.zip"
    zip_path = os.path.join(temp_dir, zip_filename)
    
    try:
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract embedded images or render pages
        reader = PdfReader(input_path)
        img_count = 0
        with zipfile.ZipFile(zip_path, 'w') as zip_file:
            for page_num, page in enumerate(reader.pages):
                for img_num, image_file in enumerate(page.images):
                    img_count += 1
                    img_name = f"sayfa_{page_num+1}_gorsel_{img_num+1}_{image_file.name}"
                    zip_file.writestr(img_name, image_file.data)
            
            # Fallback: render text pages onto high-res canvas
            if img_count == 0:
                from PIL import ImageDraw
                import io
                for page_num, page in enumerate(reader.pages):
                    img = Image.new("RGB", (1240, 1754), (255, 255, 255))
                    draw = ImageDraw.Draw(img)
                    text = page.extract_text() or "(Metin Yok)"
                    y = 60
                    for line in text.splitlines():
                        if y > 1680:
                            break
                        draw.text((60, y), line[:110], fill=(15, 23, 42))
                        y += 26
                    buf = io.BytesIO()
                    save_fmt = "PNG" if img_format == "png" else "JPEG"
                    img.save(buf, format=save_fmt)
                    img_count += 1
                    zip_file.writestr(f"sayfa_{page_num+1}.{img_format}", buf.getvalue())
        
        background_tasks.add_task(cleanup_temp_dir, temp_dir)
        
        return FileResponse(
            path=zip_path,
            filename=zip_filename,
            media_type="application/zip"
        )
    except Exception as e:
        logger.error(f"PDF to Image conversion failed: {str(e)}")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

# 4. Image to PDF
@app.post("/api/image-to-pdf")
async def image_to_pdf(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...)
):
    if not files:
        raise HTTPException(status_code=400, detail="At least one image file must be uploaded.")
    
    temp_dir = create_request_temp_dir()
    pdf_path = os.path.join(temp_dir, "converted_images.pdf")
    
    try:
        pil_images = []
        for file in files:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in [".jpg", ".jpeg", ".png", ".webp", ".bmp"]:
                logger.warning(f"Skipping unsupported image format: {file.filename}")
                continue
            
            file_path = os.path.join(temp_dir, file.filename)
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            img = Image.open(file_path)
            # Convert to RGB mode (PDF standard)
            if img.mode != "RGB":
                img = img.convert("RGB")
            pil_images.append(img)
        
        if not pil_images:
            raise HTTPException(status_code=400, detail="No valid images were uploaded.")
        
        # Save all images into a single PDF
        pil_images[0].save(
            pdf_path,
            save_all=True,
            append_images=pil_images[1:]
        )
        
        background_tasks.add_task(cleanup_temp_dir, temp_dir)
        
        return FileResponse(
            path=pdf_path,
            filename="images_combined.pdf",
            media_type="application/pdf"
        )
    except Exception as e:
        logger.error(f"Image to PDF conversion failed: {str(e)}")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

# 5. PDF Merger & Encrypter
@app.post("/api/pdf-merge-encrypt")
async def pdf_merge_encrypt(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    password: str = Form(None)
):
    if not files:
        raise HTTPException(status_code=400, detail="At least one PDF file is required.")
    
    temp_dir = create_request_temp_dir()
    merged_path = os.path.join(temp_dir, "merged.pdf")
    
    try:
        merger = PdfWriter()
        pdf_count = 0
        
        for file in files:
            if not file.filename.lower().endswith(".pdf"):
                logger.warning(f"Skipping non-PDF file: {file.filename}")
                continue
            
            file_path = os.path.join(temp_dir, f"file_{pdf_count}.pdf")
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            merger.append(file_path)
            pdf_count += 1
            
        if pdf_count == 0:
            raise HTTPException(status_code=400, detail="No valid PDF files were uploaded.")
        
        # Apply password encryption if requested
        if password:
            merger.encrypt(password)
            
        merger.write(merged_path)
        merger.close()
        
        background_tasks.add_task(cleanup_temp_dir, temp_dir)
        
        output_filename = "secured_document.pdf" if password else "merged_document.pdf"
        return FileResponse(
            path=merged_path,
            filename=output_filename,
            media_type="application/pdf"
        )
    except Exception as e:
        logger.error(f"PDF merge/encrypt operation failed: {str(e)}")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f"Operation error: {str(e)}")

# 6. Image Converter (Format conversion)
@app.post("/api/convert-image")
async def convert_image(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    target_format: str = Form(...)  # png, jpg, webp
):
    target_format = target_format.lower()
    if target_format not in ["png", "jpg", "jpeg", "webp"]:
        raise HTTPException(status_code=400, detail="Unsupported target image format.")
    
    temp_dir = create_request_temp_dir()
    input_path = os.path.join(temp_dir, file.filename)
    
    ext = "jpg" if target_format in ["jpg", "jpeg"] else target_format
    output_filename = f"converted.{ext}"
    output_path = os.path.join(temp_dir, output_filename)
    
    try:
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        img = Image.open(input_path)
        save_format = "PNG"
        if target_format in ["jpg", "jpeg"]:
            save_format = "JPEG"
        elif target_format == "webp":
            save_format = "WEBP"
            
        # If target is JPEG, flatten transparency to white background
        if save_format == "JPEG" and img.mode in ("RGBA", "LA"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3] if img.mode == "RGBA" else img.split()[1])
            bg.save(output_path, save_format)
        else:
            img.save(output_path, save_format)
            
        background_tasks.add_task(cleanup_temp_dir, temp_dir)
        
        orig_name_no_ext = os.path.splitext(file.filename)[0]
        return FileResponse(
            path=output_path,
            filename=f"{orig_name_no_ext}.{ext}",
            media_type=f"image/{ext}"
        )
    except Exception as e:
        logger.error(f"Image conversion failed: {str(e)}")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise HTTPException(status_code=500, detail=f"Conversion error: {str(e)}")

@app.post("/convert/compress-pdf")
async def compress_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Lütfen bir PDF dosyası yükleyin.")
    temp_dir = tempfile.mkdtemp()
    try:
        input_path = os.path.join(temp_dir, "input.pdf")
        output_path = os.path.join(temp_dir, "compressed.pdf")
        with open(input_path, "wb") as f:
            f.write(await file.read())
            
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
            
        with open(output_path, "wb") as f:
            writer.write(f)
            
        out_name = f"compressed_{file.filename}"
        return FileResponse(output_path, filename=out_name, media_type="application/pdf")
    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Sıkıştırma hatası: {str(e)}")

@app.post("/convert/watermark-pdf")
async def watermark_pdf(file: UploadFile = File(...), watermark_text: str = Form("Game of PDF")):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Lütfen bir PDF dosyası yükleyin.")
    temp_dir = tempfile.mkdtemp()
    try:
        input_path = os.path.join(temp_dir, "input.pdf")
        watermark_pdf_path = os.path.join(temp_dir, "watermark.pdf")
        output_path = os.path.join(temp_dir, "watermarked.pdf")
        with open(input_path, "wb") as f:
            f.write(await file.read())
            
        from reportlab.pdfgen import canvas
        from reportlab.lib.colors import Color
        
        reader = PdfReader(input_path)
        writer = PdfWriter()
        
        for page in reader.pages:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            
            c = canvas.Canvas(watermark_pdf_path, pagesize=(width, height))
            c.setFont("Helvetica-Bold", 36)
            c.setFillColor(Color(0.5, 0.5, 0.5, alpha=0.3))
            c.saveState()
            c.translate(width / 2, height / 2)
            c.rotate(45)
            c.drawCentredString(0, 0, watermark_text)
            c.restoreState()
            c.save()
            
            wm_reader = PdfReader(watermark_pdf_path)
            page.merge_page(wm_reader.pages[0])
            writer.add_page(page)
            
        with open(output_path, "wb") as f:
            writer.write(f)
            
        out_name = f"watermarked_{file.filename}"
        return FileResponse(output_path, filename=out_name, media_type="application/pdf")
    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=f"Filigran hatası: {str(e)}")

@app.get("/healthz")
async def health_check():
    return {
        "status": "ok",
        "startup_error": startup_error,
        "pil_loaded": Image is not None,
        "docx_loaded": Document is not None,
        "pypdf_loaded": PdfReader is not None
    }

@app.get("/googlea587dd58ff1846df.html")
async def serve_google_verification_file():
    return HTMLResponse("google-site-verification: googlea587dd58ff1846df.html", media_type="text/html")

@app.get("/robots.txt")
async def serve_robots_file():
    content = "User-agent: *\nAllow: /\n\nSitemap: https://gameofpdf.com/sitemap.xml"
    return HTMLResponse(content, media_type="text/plain")

@app.get("/ads.txt")
async def serve_ads_file():
    content = "google.com, pub-8999159765450234, DIRECT, f08c47fec0942fa0"
    return HTMLResponse(content, media_type="text/plain")

@app.get("/sitemap.xml")
async def serve_sitemap_file():
    static_sitemap = os.path.join(os.path.dirname(__file__), "static", "sitemap.xml")
    if os.path.exists(static_sitemap):
        return FileResponse(static_sitemap, media_type="application/xml")
    xml_content = '<?xml version="1.0" encoding="UTF-8"?><urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"><url><loc>https://gameofpdf.com/</loc></url></urlset>'
    return HTMLResponse(xml_content, media_type="application/xml")





